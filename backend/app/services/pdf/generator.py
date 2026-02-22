"""
PDF CV Generator — Olive template.
Renders an orange header bar (full page width) + two-column body layout that
matches the live preview shown in the CV editor.
"""
from __future__ import annotations

import html
import io
import os
from typing import Any

from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    HRFlowable,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)

# ---------------------------------------------------------------------------
# Font registration — prefer Arial (Unicode TTF) over Helvetica (latin-1)
# ---------------------------------------------------------------------------
_FONT_REG = "Helvetica"
_FONT_BOLD = "Helvetica-Bold"

_win_fonts = os.path.join(os.environ.get("SystemRoot", "C:\\Windows"), "Fonts")
try:
    _f_reg = os.path.join(_win_fonts, "arial.ttf")
    _f_bold = os.path.join(_win_fonts, "arialbd.ttf")
    if os.path.exists(_f_reg):
        pdfmetrics.registerFont(TTFont("ArialUni", _f_reg))
        _FONT_REG = "ArialUni"
    if os.path.exists(_f_bold):
        pdfmetrics.registerFont(TTFont("ArialUni-Bold", _f_bold))
        _FONT_BOLD = "ArialUni-Bold"
except Exception:
    pass  # Fall back to Helvetica silently

# True when a TTF Unicode font is loaded — no latin-1 restriction applies
_USE_UNICODE = _FONT_REG != "Helvetica"

# ---------------------------------------------------------------------------
# Colour palette  (mirrors preview-panel.tsx Olive template)
# ---------------------------------------------------------------------------
ORANGE = colors.HexColor("#dda15e")       # header bar / accent
DARK_GREEN = colors.HexColor("#283618")   # name, headings, section titles
MED_GREEN = colors.HexColor("#606c38")    # secondary text
WHITE = colors.white
BLACK = colors.HexColor("#1a1a1a")
DIVIDER_C = colors.HexColor("#c8d5b9")   # thin divider lines

# ---------------------------------------------------------------------------
# Page geometry
# ---------------------------------------------------------------------------
PAGE_W, PAGE_H = A4               # 595 × 842 pt
MARGIN_H: float = 18 * mm         # left/right margin
MARGIN_BOT: float = 14 * mm       # bottom margin
HEADER_H: float = 62              # orange bar height (pt)
HEADER_GAP: float = 10            # gap between header bar and content

# Body frame dimensions
CONTENT_W: float = PAGE_W - 2 * MARGIN_H
CONTENT_H: float = PAGE_H - HEADER_H - HEADER_GAP - MARGIN_BOT

SIDEBAR_W: float = 128            # right sidebar width (pt ≈ 45 mm)
SIDEBAR_GAP: float = 10           # gap between main and sidebar
MAIN_W: float = CONTENT_W - SIDEBAR_W - SIDEBAR_GAP


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

# Map common Unicode typography → ASCII/latin-1 equivalents
_UNICODE_MAP = str.maketrans({
    "\u2013": "-",      # en dash
    "\u2014": " - ",    # em dash
    "\u2015": "-",      # horizontal bar
    "\u2018": "'",      # left single quote
    "\u2019": "'",      # right single quote
    "\u201a": ",",      # single low-9 quotation mark
    "\u201c": '"',      # left double quote
    "\u201d": '"',      # right double quote
    "\u2022": "-",      # bullet
    "\u2026": "...",    # ellipsis
    "\u00a0": " ",      # non-breaking space
    "\u2011": "-",      # non-breaking hyphen
    "\u00b7": ".",      # middle dot
    "\u2192": "->",     # right arrow
    "\u2019": "'",      # duplicate safety
})


def _safe(val: Any, fallback: str = "") -> str:
    """Return a string safe for ReportLab (HTML-escaped).
    When ArialUni TTF is loaded, full Unicode is supported — no latin-1 stripping.
    Falls back to ASCII-safe encoding only when using Helvetica (latin-1 font).
    """
    if val is None:
        return fallback
    text = str(val).strip()
    if not _USE_UNICODE:
        # Helvetica only: replace known Unicode chars, then strip the rest
        text = text.translate(_UNICODE_MAP)
        text = text.encode("latin-1", errors="ignore").decode("latin-1")
    return html.escape(text)


def _styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        # Main column section titles (e.g. "WORK EXPERIENCE")
        "section_title": ParagraphStyle(
            "section_title",
            parent=base["Normal"],
            fontSize=8,
            leading=11,
            textColor=DARK_GREEN,
            fontName=_FONT_BOLD,
            spaceBefore=8,
            spaceAfter=1,
        ),
        # Normal body text
        "body": ParagraphStyle(
            "body",
            parent=base["Normal"],
            fontSize=8.5,
            leading=12.5,
            textColor=BLACK,
        ),
        # Right-aligned muted date / meta text
        "small_right": ParagraphStyle(
            "small_right",
            parent=base["Normal"],
            fontSize=7.5,
            leading=11,
            textColor=MED_GREEN,
            alignment=TA_RIGHT,
        ),
        # Bullet lines (indented)
        "bullet": ParagraphStyle(
            "bullet",
            parent=base["Normal"],
            fontSize=8,
            leading=11.5,
            textColor=BLACK,
            leftIndent=10,
            spaceAfter=1,
        ),
        # Sidebar section title
        "sidebar_section": ParagraphStyle(
            "sidebar_section",
            parent=base["Normal"],
            fontSize=7.5,
            leading=10,
            textColor=DARK_GREEN,
            fontName=_FONT_BOLD,
            spaceBefore=8,
            spaceAfter=1,
        ),
        # Sidebar body text
        "sidebar_body": ParagraphStyle(
            "sidebar_body",
            parent=base["Normal"],
            fontSize=7.5,
            leading=11,
            textColor=BLACK,
        ),
        # Sidebar skill dots
        "sidebar_skill": ParagraphStyle(
            "sidebar_skill",
            parent=base["Normal"],
            fontSize=7.5,
            leading=10,
            textColor=DARK_GREEN,
            leftIndent=6,
        ),
    }


def _hr() -> HRFlowable:
    return HRFlowable(width="100%", thickness=0.4, color=DIVIDER_C, spaceAfter=3)


def _section(items: list, title: str, s: dict) -> None:
    """Append a section header + divider to items list."""
    items.append(Paragraph(title.upper(), s["section_title"]))
    items.append(_hr())


def _sidebar_section(items: list, title: str, s: dict) -> None:
    items.append(Paragraph(title.upper(), s["sidebar_section"]))
    items.append(_hr())


def _exp_row(left: str, right: str, s: dict) -> Table:
    """Job title / school line with right-aligned date."""
    t = Table(
        [[Paragraph(f"<b>{left}</b>", s["body"]), Paragraph(right, s["small_right"])]],
        colWidths=[MAIN_W * 0.68, MAIN_W * 0.32],
    )
    t.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
    ]))
    return t


# ---------------------------------------------------------------------------
# Main generator
# ---------------------------------------------------------------------------

def _clean(obj: Any) -> Any:
    """Recursively sanitize all strings in a nested structure.
    When ArialUni TTF is active, strings are left as-is (full Unicode support).
    Otherwise, maps common Unicode chars to ASCII equivalents and strips the rest.
    Non-string values (datetime, int, bool, None…) are always left untouched.
    """
    if isinstance(obj, str):
        if not _USE_UNICODE:
            obj = obj.translate(_UNICODE_MAP)
            return obj.encode("latin-1", errors="ignore").decode("latin-1")
        return obj
    if isinstance(obj, dict):
        return {k: _clean(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [_clean(item) for item in obj]
    return obj


async def generate_cv_pdf(cv: dict) -> bytes:
    """
    Build a styled PDF from a CV Firestore document and return raw bytes.
    Visual style: Olive template (orange header bar, two-column body).
    """
    # ── Sanitize ALL string values recursively ───────────────────────────────
    # Walks every dict/list in the CV, replaces non-latin-1 chars in strings.
    # Non-string types (Firestore Timestamps, ints…) are left untouched.
    cv = _clean(cv)

    content: dict = cv.get("content") or {}
    doc_title: str = _safe(cv.get("title"), "Curriculum Vitae")

    # --- Header data --------------------------------------------------------
    contact: dict = content.get("contact_info") or {}
    name = _safe(contact.get("name")) or doc_title
    first_exp = next(iter(content.get("experience") or []), {})
    job_title_str = _safe(first_exp.get("job_title"))

    # --- Canvas callback: draws orange header on every page -----------------
    def _on_page(canvas, _doc):
        canvas.saveState()
        # Full-width orange bar
        canvas.setFillColor(ORANGE)
        canvas.rect(0, PAGE_H - HEADER_H, PAGE_W, HEADER_H, fill=1, stroke=0)
        # Name — white bold
        canvas.setFillColor(WHITE)
        canvas.setFont(_FONT_BOLD, 18)
        canvas.drawString(MARGIN_H, PAGE_H - 36, name[:65])
        # Job title — dark text on orange
        if job_title_str:
            canvas.setFillColor(DARK_GREEN)
            canvas.setFont(_FONT_REG, 9)
            canvas.drawString(MARGIN_H, PAGE_H - 51, job_title_str[:90])
        canvas.restoreState()

    # --- Build doc ----------------------------------------------------------
    buf = io.BytesIO()

    body_frame = Frame(
        MARGIN_H,
        MARGIN_BOT,
        CONTENT_W,
        CONTENT_H,
        leftPadding=0,
        rightPadding=0,
        topPadding=0,
        bottomPadding=0,
        id="body",
    )

    doc = BaseDocTemplate(
        buf,
        pagesize=A4,
        title=doc_title,
        author="CVFlow",
    )
    doc.addPageTemplates([
        PageTemplate(id="main", frames=[body_frame], onPage=_on_page)
    ])

    s = _styles()

    # ========================================================================
    # MAIN COLUMN
    # ========================================================================
    main: list = []

    # Summary ----------------------------------------------------------------
    summary = _safe(content.get("summary"))
    if summary:
        _section(main, "Professional Summary", s)
        main.append(Paragraph(summary, s["body"]))
        main.append(Spacer(1, 4))

    # Experience -------------------------------------------------------------
    experience: list[dict] = content.get("experience") or []
    if experience:
        _section(main, "Work Experience", s)
        for i, job in enumerate(experience):
            jt = _safe(job.get("job_title"))
            co = _safe(job.get("company"))
            loc = _safe(job.get("location"))
            start = _safe(job.get("start_date"))
            end = "Present" if job.get("current") else _safe(job.get("end_date"))
            date_range = " - ".join(filter(None, [start, end]))
            left_text = " | ".join(filter(None, [jt, co, loc]))

            if date_range:
                main.append(_exp_row(left_text, date_range, s))
            else:
                main.append(Paragraph(f"<b>{left_text}</b>", s["body"]))

            for bullet in (job.get("bullets") or []):
                main.append(Paragraph(f"• {_safe(bullet)}", s["bullet"]))

            if i < len(experience) - 1:
                main.append(Spacer(1, 5))
        main.append(Spacer(1, 2))

    # Education --------------------------------------------------------------
    education: list[dict] = content.get("education") or []
    if education:
        _section(main, "Education", s)
        for i, edu in enumerate(education):
            school = _safe(edu.get("school"))
            degree = _safe(edu.get("degree"))
            field = _safe(edu.get("field"))
            grad = _safe(edu.get("graduation_date"))
            gpa = _safe(edu.get("gpa"))

            degree_line = " in ".join(filter(None, [degree, field]))
            left = " | ".join(filter(None, [school, degree_line]))
            right_parts = list(filter(None, [grad, f"GPA {gpa}" if gpa else ""]))
            right = " | ".join(right_parts)

            if right:
                main.append(_exp_row(left, right, s))
            else:
                main.append(Paragraph(f"<b>{left}</b>", s["body"]))

            if i < len(education) - 1:
                main.append(Spacer(1, 4))
        main.append(Spacer(1, 2))

    # Certifications ---------------------------------------------------------
    certifications: list[str] = content.get("certifications") or []
    if certifications:
        _section(main, "Certifications", s)
        for cert in certifications:
            main.append(Paragraph(f"• {_safe(cert)}", s["bullet"]))
        main.append(Spacer(1, 2))

    # Projects ---------------------------------------------------------------
    projects: list[dict] = content.get("projects") or []
    if projects:
        _section(main, "Projects", s)
        for i, proj in enumerate(projects):
            proj_name = _safe(proj.get("name") or proj.get("title", ""))
            proj_desc = _safe(proj.get("description", ""))
            proj_url = _safe(proj.get("url", ""))
            if proj_name:
                main.append(Paragraph(f"<b>{proj_name}</b>", s["body"]))
            if proj_desc:
                main.append(Paragraph(proj_desc, s["body"]))
            if proj_url:
                main.append(Paragraph(proj_url, s["bullet"]))
            if i < len(projects) - 1:
                main.append(Spacer(1, 4))

    # ========================================================================
    # SIDEBAR
    # ========================================================================
    sidebar: list = []

    # Contact ----------------------------------------------------------------
    _sidebar_section(sidebar, "Contact", s)
    for val in filter(None, [
        _safe(contact.get("email")),
        _safe(contact.get("phone")),
        _safe(contact.get("location")),
        _safe(contact.get("linkedin")),
        _safe(contact.get("website")),
    ]):
        sidebar.append(Paragraph(val, s["sidebar_body"]))
        sidebar.append(Spacer(1, 1.5))

    # Skills -----------------------------------------------------------------
    skills: list[str] = content.get("skills") or []
    if skills:
        _sidebar_section(sidebar, "Skills", s)
        for skill in skills:
            sidebar.append(Paragraph(f"• {_safe(skill)}", s["sidebar_skill"]))
            sidebar.append(Spacer(1, 1))

    # Languages --------------------------------------------------------------
    languages: list[str] = content.get("languages") or []
    if languages:
        _sidebar_section(sidebar, "Languages", s)
        for lang in languages:
            sidebar.append(Paragraph(_safe(lang), s["sidebar_body"]))
            sidebar.append(Spacer(1, 1.5))

    # ========================================================================
    # TWO-COLUMN BODY TABLE
    # ========================================================================
    # If main is empty (fresh CV), show a placeholder
    if not main:
        main.append(Spacer(1, 4))

    body_table = Table(
        [[main, sidebar]],
        colWidths=[MAIN_W, SIDEBAR_W],
    )
    body_table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        # Main column padding
        ("LEFTPADDING", (0, 0), (0, 0), 0),
        ("RIGHTPADDING", (0, 0), (0, 0), 4),
        # Sidebar padding
        ("LEFTPADDING", (1, 0), (1, 0), SIDEBAR_GAP),
        ("RIGHTPADDING", (1, 0), (1, 0), 0),
        # Remove default vertical padding
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        # Thin vertical divider between columns
        ("LINEBEFORE", (1, 0), (1, -1), 0.4, DIVIDER_C),
    ]))

    doc.build([body_table])
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Cover Letter PDF generator
# ---------------------------------------------------------------------------

async def generate_cover_letter_pdf(
    text: str,
    tone: str = "professional",
    letter_format: str = "us",
) -> bytes:
    """Render a cover letter as a clean, professional PDF and return bytes."""
    base = getSampleStyleSheet()

    CL_BODY = ParagraphStyle(
        "cl_body",
        parent=base["Normal"],
        fontName=_FONT_REG,
        fontSize=10.5,
        leading=16,
        textColor=colors.HexColor("#1a1a1a"),
        spaceAfter=10,
    )

    buf = io.BytesIO()
    doc = BaseDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=25 * mm,
        rightMargin=25 * mm,
        topMargin=22 * mm,
        bottomMargin=22 * mm,
        title="Cover Letter",
        author="CVFlow",
    )

    frame = Frame(
        doc.leftMargin, doc.bottomMargin,
        doc.width, doc.height,
        leftPadding=0, rightPadding=0,
        topPadding=0, bottomPadding=0,
    )
    doc.addPageTemplates([PageTemplate(id="main", frames=[frame])])

    story = []
    for para in text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        if _USE_UNICODE:
            safe = html.escape(para)
        else:
            safe = para.translate(_UNICODE_MAP)
            safe = safe.encode("latin-1", errors="ignore").decode("latin-1")
            safe = html.escape(safe)
        story.append(Paragraph(safe, CL_BODY))

    if not story:
        story.append(Paragraph("(empty cover letter)", CL_BODY))

    doc.build(story)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Cover Letter DOCX generator
# ---------------------------------------------------------------------------

async def generate_cover_letter_docx(
    text: str,
    tone: str = "professional",
    letter_format: str = "us",
) -> bytes:
    """Render a cover letter as a .docx file and return bytes."""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins (2.5 cm each side)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    for para in text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        p = doc.add_paragraph()
        run = p.add_run(para)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = Pt(15)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
